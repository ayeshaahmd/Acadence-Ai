from flask import Blueprint, request, jsonify
from app import db
from app.models import ResearchPaper, ResearchSource
from app.services.generation_service import PaperGenerationService
from app.config import config
import logging
import io
import base64

# Import exporting libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib import colors

logger = logging.getLogger(__name__)

bp = Blueprint('papers', __name__, url_prefix='/api/papers')
generation_service = PaperGenerationService(config)

# --- EXPORT HELPER FUNCTIONS ---

def generate_docx_bytes(title, content):
    """Generate structured DOCX file from paper content"""
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup (Times New Roman, academic standard)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Paper Title (Large, Bold, Centered)
    t_p = doc.add_paragraph()
    t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = t_p.add_run(title)
    t_run.font.size = Pt(18)
    t_run.bold = True
    
    doc.add_paragraph()  # Spacing
    
    lines = content.split('\n')
    in_table = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Already handled title
            continue
            
        elif line.startswith('## '):
            in_table = False
            heading_text = line.replace('## ', '')
            h_p = doc.add_paragraph()
            h_run = h_p.add_run(heading_text)
            h_run.bold = True
            h_run.font.size = Pt(14)
            h_p.paragraph_format.space_before = Pt(18)
            h_p.paragraph_format.space_after = Pt(6)
            
        elif line.startswith('### '):
            in_table = False
            heading_text = line.replace('### ', '')
            h_p = doc.add_paragraph()
            h_run = h_p.add_run(heading_text)
            h_run.bold = True
            h_run.font.italic = True
            h_run.font.size = Pt(12)
            h_p.paragraph_format.space_before = Pt(12)
            h_p.paragraph_format.space_after = Pt(4)
            
        elif line.startswith('|') and (line.endswith('|') or line.count('|') > 1):
            # Simple markdown table parser for DOCX
            cols = [col.strip() for col in line.split('|')][1:-1]
            if cols and not all(c.startswith('-') for c in cols):
                if not in_table:
                    # Create new table in Word
                    in_table = True
                    word_table = doc.add_table(rows=0, cols=len(cols))
                    word_table.style = 'Table Grid'
                
                row_cells = word_table.add_row().cells
                for idx, col_val in enumerate(cols):
                    row_cells[idx].text = col_val
        else:
            in_table = False
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5  # 1.5 line spacing
            p.paragraph_format.space_after = Pt(6)
            p.add_run(line)
            
    # Save doc to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


def generate_pdf_bytes(title, content):
    """Generate professional typeset academic PDF using reportlab"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    # Custom academic styles
    title_style = ParagraphStyle(
        'AcademicTitle',
        parent=styles['Heading1'],
        fontName='Times-Bold',
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    heading_style = ParagraphStyle(
        'AcademicHeading',
        parent=styles['Heading2'],
        fontName='Times-Bold',
        fontSize=14,
        leading=18,
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )
    
    subheading_style = ParagraphStyle(
        'AcademicSubheading',
        parent=styles['Heading3'],
        fontName='Times-BoldItalic',
        fontSize=12,
        leading=16,
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'AcademicBody',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=8
    )
    
    table_text_style = ParagraphStyle(
        'TableText',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=9,
        leading=11
    )
    
    story = []
    
    # Title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 15))
    
    lines = content.split('\n')
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            continue
            
        elif line.startswith('## '):
            if in_table and table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.black),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
                    ('BOTTOMPADDING', (0,0), (-1,0), 6),
                    ('BACKGROUND', (0,1), (-1,-1), colors.whitesmoke),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ]))
                story.append(t)
                story.append(Spacer(1, 10))
                in_table = False
                table_data = []
                
            heading_text = line.replace('## ', '')
            if heading_text.lower().startswith('references'):
                story.append(PageBreak())  # Start references on new page
            story.append(Paragraph(heading_text, heading_style))
            
        elif line.startswith('### '):
            subheading_text = line.replace('### ', '')
            story.append(Paragraph(subheading_text, subheading_style))
            
        elif line.startswith('|') and (line.endswith('|') or line.count('|') > 1):
            in_table = True
            cols = [col.strip() for col in line.split('|')][1:-1]
            if cols and not all(c.startswith('-') for c in cols):
                row_paras = [Paragraph(col, table_text_style) for col in cols]
                table_data.append(row_paras)
        else:
            if in_table and table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.black),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
                    ('BOTTOMPADDING', (0,0), (-1,0), 6),
                    ('BACKGROUND', (0,1), (-1,-1), colors.whitesmoke),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ]))
                story.append(t)
                story.append(Spacer(1, 10))
                in_table = False
                table_data = []
                
            story.append(Paragraph(line, body_style))
            
    if in_table and table_data:
        t = Table(table_data)
        t.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
        ]))
        story.append(t)
        
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_latex_content(title, abstract, content):
    """Generate high-quality compilable LaTeX source code"""
    latex = r"""\documentclass[12pt,journal]{article}
\usepackage[utf8]{inputenc}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{booktabs}
\usepackage{hyperref}
\usepackage{graphicx}
\usepackage{geometry}
\geometry{letterpaper, margin=1in}

\begin{document}

\title{""" + title + r"""}
\author{Research Paper Assistant - PhD Suite}
\date{\today}
\maketitle

\begin{abstract}
""" + abstract + r"""
\end{abstract}

"""
    lines = content.split('\n')
    in_references = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            continue
        elif line.startswith('## Abstract'):
            continue
        elif line.startswith('## References'):
            in_references = True
            latex += "\n\\begin{thebibliography}{99}\n"
            continue
        elif line.startswith('## '):
            heading = line.replace('## ', '')
            latex += f"\n\\section{{{heading}}}\n"
        elif line.startswith('### '):
            subheading = line.replace('### ', '')
            latex += f"\n\\subsection{{{subheading}}}\n"
        else:
            if in_references:
                ref_text = line
                if '.' in line[:4]:
                    ref_text = line.split('.', 1)[1].strip()
                cite_key = ref_text[:12].replace(' ', '').replace(',', '').replace('.', '')
                latex += f"\\bibitem{{{cite_key}}} {ref_text}\n"
            else:
                # Escape standard LaTeX special characters safely
                line_escaped = line.replace('%', '\\%').replace('_', '\\_').replace('&', '\\&').replace('$', '\\$')
                # But restore Math blocks $$ back
                line_escaped = line_escaped.replace('\\$\\$', '$$')
                latex += f"{line_escaped}\n\n"
                
    if in_references:
        latex += "\\end{thebibliography}\n"
        
    latex += "\n\\end{document}\n"
    return latex


# --- ENDPOINTS ---

@bp.route('', methods=['POST'])
def generate_paper():
    """Generate a new research paper"""
    try:
        data = request.get_json()
        
        # Validate input
        if not data or 'topic' not in data:
            return jsonify({'error': 'Topic is required'}), 400
        
        topic = data.get('topic')
        word_count = int(data.get('word_count', config.DEFAULT_WORD_COUNT))
        citation_format = data.get('citation_format', config.DEFAULT_CITATION_FORMAT)
        complexity_level = data.get('complexity_level', 'master')
        domain = data.get('domain', 'general')
        
        # Validate ranges
        word_count = max(config.MIN_WORD_COUNT, min(word_count, config.MAX_WORD_COUNT))
        
        if citation_format not in config.CITATION_FORMATS:
            return jsonify({'error': f'Invalid citation format. Use one of: {", ".join(config.CITATION_FORMATS)}'}), 400
        
        # Generate paper
        result = generation_service.generate_complete_paper(
            topic, word_count, citation_format, complexity_level, domain
        )
        
        return jsonify({
            'success': True,
            'data': result
        }), 201
    
    except ValueError as e:
        logger.error(f"Validation error: {e}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Paper generation error: {e}")
        return jsonify({'error': 'Failed to generate paper', 'details': str(e)}), 500


@bp.route('', methods=['GET'])
def list_papers():
    """List all generated papers"""
    try:
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 10))
        
        skip = (page - 1) * per_page
        papers_qs = ResearchPaper.objects.order_by('-created_at')
        total = papers_qs.count()
        papers = papers_qs.skip(skip).limit(per_page)
        
        return jsonify({
            'success': True,
            'data': [paper.to_dict() for paper in papers],
            'total': total,
            'pages': (total + per_page - 1) // per_page,
            'current_page': page
        }), 200
    except Exception as e:
        logger.error(f"List papers error: {e}")
        return jsonify({'error': 'Failed to list papers'}), 500


@bp.route('/<string:paper_id>', methods=['GET'])
def get_paper(paper_id):
    """Get a specific paper"""
    try:
        paper = ResearchPaper.objects(id=paper_id).first()
        
        if not paper:
            return jsonify({'error': 'Paper not found'}), 404
        
        return jsonify({
            'success': True,
            'data': {
                **paper.to_dict(),
                'content': paper.content
            }
        }), 200
    except Exception as e:
        logger.error(f"Get paper error: {e}")
        return jsonify({'error': 'Failed to retrieve paper'}), 500


@bp.route('/<string:paper_id>', methods=['DELETE'])
def delete_paper(paper_id):
    """Delete a paper"""
    try:
        paper = ResearchPaper.objects(id=paper_id).first()
        
        if not paper:
            return jsonify({'error': 'Paper not found'}), 404
        
        paper.delete()
        
        return jsonify({
            'success': True,
            'message': 'Paper deleted successfully'
        }), 200
    except Exception as e:
        logger.error(f"Delete paper error: {e}")
        return jsonify({'error': 'Failed to delete paper'}), 500


@bp.route('/<string:paper_id>/export', methods=['GET'])
def export_paper(paper_id):
    """Export paper as file in multiple academic formats"""
    try:
        paper = ResearchPaper.objects(id=paper_id).first()
        
        if not paper:
            return jsonify({'error': 'Paper not found'}), 404
        
        file_format = request.args.get('format', 'txt').lower()
        
        if file_format == 'pdf':
            pdf_bytes = generate_pdf_bytes(paper.title, paper.content)
            base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
            return jsonify({
                'success': True,
                'filename': f"{paper.title.replace(' ', '_')}.pdf",
                'content': base64_pdf,
                'is_binary': True
            }), 200
            
        elif file_format == 'docx':
            docx_bytes = generate_docx_bytes(paper.title, paper.content)
            base64_docx = base64.b64encode(docx_bytes).decode('utf-8')
            return jsonify({
                'success': True,
                'filename': f"{paper.title.replace(' ', '_')}.docx",
                'content': base64_docx,
                'is_binary': True
            }), 200
            
        elif file_format == 'latex' or file_format == 'tex':
            latex_content = generate_latex_content(paper.title, paper.abstract, paper.content)
            return jsonify({
                'success': True,
                'filename': f"{paper.title.replace(' ', '_')}.tex",
                'content': latex_content,
                'is_binary': False
            }), 200
            
        else:
            # Plain text export
            return jsonify({
                'success': True,
                'filename': f"{paper.title.replace(' ', '_')}.txt",
                'content': paper.content,
                'is_binary': False
            }), 200
            
    except Exception as e:
        logger.error(f"Export paper error: {e}")
        return jsonify({'error': 'Failed to export paper', 'details': str(e)}), 500
